from docx import Document
from .base import DocumentReader


class DocxReader:
    def read(self, path: str) -> str:
        document = Document(path)
        parts: list[str] = []

        # --- Paragraphs ---
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # --- Tables ---
        for table in document.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if row_text:
                    parts.append(" | ".join(row_text))

        return "\n".join(parts)
